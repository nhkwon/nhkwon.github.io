from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "supabase-board-guide.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, size=10, bold=False, color=None):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    set_font(run, 9.5, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 13, True, (31, 78, 121))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_font(run, 10.5)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.8)
    run.font.color.rgb = RGBColor(36, 41, 47)


def add_step_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, width in enumerate([0.65, 1.7, 4.25]):
        table.columns[i].width = Inches(width)

    for i, header in enumerate(["순서", "항목", "내용"]):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "1F4E79")
        set_cell_text(cell, header, True, (255, 255, 255))

    for no, item, detail in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], no, True, (31, 78, 121))
        set_cell_text(cells[1], item, True)
        set_cell_text(cells[2], detail)

    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

doc.styles["Normal"].font.name = "Malgun Gothic"
doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
doc.styles["Normal"].font.size = Pt(10.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Supabase를 활용한 게시판 연결 가이드")
set_font(run, 20, True, (31, 78, 121))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Vercel + Supabase + Contact 게시판 설정 절차")
set_font(run, 11, False, (89, 99, 110))

add_heading(doc, "개요")
add_body(
    doc,
    "이 문서는 홈페이지 Contact 페이지의 게시판을 Supabase 데이터베이스와 연결하는 전체 절차를 정리한 문서입니다. "
    "브라우저가 Supabase 키를 직접 쓰지 않고, Vercel API가 중간에서 Supabase에 저장하는 구조입니다.",
)

add_heading(doc, "전체 흐름")
add_step_table(
    doc,
    [
        (1, "사용자 입력", "Contact 페이지에서 이름, 이메일, 제목, 내용, 첨부파일을 입력합니다."),
        (2, "Vercel API 호출", "브라우저가 /api/contact-board 서버 함수로 등록 요청을 보냅니다."),
        (3, "Supabase 저장", "Vercel 서버 함수가 환경변수의 Supabase 키를 사용해 contact_posts 테이블에 저장합니다."),
        (4, "목록 표시", "Contact 페이지가 Supabase에서 게시글 목록을 불러와 화면에 표시합니다."),
    ],
)

add_heading(doc, "1. Supabase 프로젝트 생성")
add_body(doc, "Supabase에 로그인한 뒤 New project를 눌러 새 프로젝트를 생성합니다. 프로젝트 이름, 데이터베이스 비밀번호, region을 설정하고 생성이 완료될 때까지 기다립니다.")

add_heading(doc, "2. 게시판 테이블 생성")
add_body(doc, "Supabase 왼쪽 메뉴에서 SQL Editor를 열고 New query를 선택한 뒤 아래 SQL을 실행합니다.")
add_code(
    doc,
    """create extension if not exists pgcrypto;

create table if not exists public.contact_posts (
  id uuid primary key default gen_random_uuid(),
  name text not null,
  email text not null,
  subject text not null,
  message text not null,
  attachments jsonb not null default '[]'::jsonb,
  created_at timestamptz not null default now()
);

alter table public.contact_posts enable row level security;

create index if not exists contact_posts_created_at_idx
  on public.contact_posts (created_at desc);""",
)

add_heading(doc, "3. RLS 정책 추가")
add_body(doc, "RLS가 켜져 있으면 권한 정책이 없을 때 등록이 막힙니다. SQL Editor에서 아래 정책을 실행합니다.")
add_code(
    doc,
    """create policy "contact_posts_select_public"
  on public.contact_posts
  for select
  using (true);

create policy "contact_posts_insert_public"
  on public.contact_posts
  for insert
  with check (true);

create policy "contact_posts_delete_public"
  on public.contact_posts
  for delete
  using (true);""",
)

add_heading(doc, "4. Supabase URL 값 찾기")
add_body(doc, "Supabase에서 Project Settings → API로 이동합니다. 화면의 Project URL 값을 복사합니다.")
add_code(doc, "예시: https://xxxxxx.supabase.co")
add_body(doc, "Data API 화면에서 API URL이 https://xxxxxx.supabase.co/rest/v1/ 형태로 보일 수 있습니다. Vercel의 SUPABASE_URL에는 /rest/v1/을 제외하고 https://xxxxxx.supabase.co까지만 넣어야 합니다.")

add_heading(doc, "5. Supabase service_role key 찾기")
add_body(doc, "Supabase에서 Project Settings → API → Project API keys로 이동합니다. 여기에서 service_role 또는 secret service_role key를 복사합니다. anon public key가 아니라 service_role key를 사용해야 합니다.")

add_heading(doc, "6. Vercel 환경변수 추가")
add_body(doc, "Vercel에서 프로젝트를 선택한 뒤 Settings → Environment Variables로 이동합니다. Add Environment Variable을 눌러 아래 두 값을 추가합니다.")
add_step_table(
    doc,
    [
        ("A", "SUPABASE_URL", "Supabase Project URL 값을 넣습니다. 예: https://xxxxxx.supabase.co"),
        ("B", "SUPABASE_SERVICE_ROLE_KEY", "Supabase service_role key 값을 넣습니다. 이 값은 Sensitive로 표시되는 것이 정상입니다."),
    ],
)
add_body(doc, "Environment는 최소 Production을 체크하고, 가능하면 Production과 Preview를 모두 체크합니다.")

add_heading(doc, "7. Vercel 재배포")
add_body(doc, "환경변수는 저장만 하면 기존 배포에 바로 반영되지 않습니다. 환경변수 추가 후 반드시 재배포합니다.")
add_code(doc, "vercel deploy --prod")
add_body(doc, "또는 Vercel 화면에서 Deployments → 최신 배포의 ... → Redeploy를 선택합니다.")

add_heading(doc, "8. 작동 확인")
add_step_table(
    doc,
    [
        (1, "Contact 페이지 접속", "홈페이지의 contact.html 페이지를 엽니다."),
        (2, "게시글 등록", "이름, 이메일, 제목, 내용을 입력하고 등록 버튼을 누릅니다."),
        (3, "Supabase 확인", "Supabase Table Editor → contact_posts에서 새 행이 생성됐는지 확인합니다."),
        (4, "오류 확인", "등록 중에서 멈추거나 실패하면 /api/contact-board 응답과 RLS 정책을 확인합니다."),
    ],
)

add_heading(doc, "주의사항")
for item in [
    "SUPABASE_SERVICE_ROLE_KEY는 관리자 권한 키입니다. GitHub, HTML, JS 파일, 공개 문서에 넣으면 안 됩니다.",
    "SUPABASE_URL에는 /rest/v1/을 포함하지 않습니다.",
    "Vercel 환경변수를 추가하거나 수정한 뒤에는 항상 재배포해야 합니다.",
    "RLS 정책이 없으면 new row violates row-level security policy 오류가 발생할 수 있습니다.",
    "게시판 API는 /api/contact-board에서 동작하며, 브라우저는 Supabase에 직접 연결하지 않습니다.",
]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(item)
    set_font(run, 10.5)

add_heading(doc, "최종 체크리스트")
add_step_table(
    doc,
    [
        ("□", "contact_posts 테이블 생성", "Supabase SQL Editor에서 테이블 생성 SQL 실행"),
        ("□", "RLS 정책 생성", "select, insert, delete 정책 생성"),
        ("□", "SUPABASE_URL 추가", "Vercel Environment Variables에 Project URL 추가"),
        ("□", "SUPABASE_SERVICE_ROLE_KEY 추가", "Vercel Environment Variables에 service_role key 추가"),
        ("□", "재배포", "vercel deploy --prod 실행"),
        ("□", "등록 테스트", "Contact 페이지에서 글 등록 후 Supabase Table Editor 확인"),
    ],
)

doc.save(OUT)
print(OUT)
